"""
Page 2: Upload PI Goals Document
Upload and validate PI goals from Word documents using AI agents
"""

import streamlit as st
import sys
import time
from pathlib import Path
from typing import Dict, List, Any

# Add the app directory to Python path for imports
app_dir = Path(__file__).parent.parent
sys.path.insert(0, str(app_dir))

from components.sidebar import render_sidebar, render_page_header, render_progress_indicator, update_workflow_status
from components.file_uploader import render_file_uploader
from agents.goal_validator import GoalValidatorAgent
from utils.file_handlers import DocumentProcessor
from utils.config import get_file_upload_config, save_session_data, load_session_data, load_config
import openai
import io
from docx import Document
from docx.shared import Inches

# Page configuration
st.set_page_config(
    page_title="Upload Goals - PI Planning Dashboard",
    page_icon="📄",
    layout="wide"
)

def main():
    """Main page function"""
    
    # Render sidebar
    render_sidebar()
    
    # Page header
    render_page_header(
        title="📄 Upload PI Goals Document",
        description="Upload your PI goals document and let AI agents validate and improve them",
        step_number=2
    )
    
    # Progress indicator
    render_progress_indicator(current_step=2)
    
    # Optional workflow status check (JIRA cleanup is optional)
    workflow_status = st.session_state.get('workflow_status', {})
    if workflow_status.get('jira_wipe') != 'complete':
        st.info("""
        **💡 Optional:** You can clean your JIRA project first if you want to start fresh, but it's not required to proceed with goal analysis.
        """)
    
    # Demo document generation section
    st.markdown("### 🎭 Demo Document Generation")
    
    st.info("""
    **For demonstration purposes**, you can generate sample PI goals documents using AI.
    This helps you test the system without needing to create your own documents.
    """)
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("✅ Generate Good Example", use_container_width=True, help="Generate a document that meets all SMART criteria"):
            generate_demo_document("good")
    
    with col2:
        if st.button("❌ Generate Poor Example", use_container_width=True, help="Generate a document with common issues"):
            generate_demo_document("poor")
    
    st.markdown("---")
    
    # File upload section
    st.markdown("### 📤 Upload PI Goals Document")
    
    upload_config = get_file_upload_config()
    
    st.info("""
    **Supported formats:** Word documents (.docx, .doc), PDF files (.pdf), or plain text (.txt)
    
    **What to include in your document:**
    - High-level PI objectives
    - Business goals and outcomes
    - Success criteria and metrics
    - Key stakeholder requirements
    - Timeline and milestones
    """)
    
    # File uploader
    uploaded_file = st.file_uploader(
        "Choose your PI goals document",
        type=['docx', 'doc', 'pdf', 'txt'],
        help=f"Maximum file size: {upload_config['max_size'] // 1024 // 1024}MB"
    )
    
    if uploaded_file is not None:
        # Display file info
        st.success(f"✅ File uploaded: **{uploaded_file.name}** ({uploaded_file.size:,} bytes)")
        
        # Process the document
        process_document(uploaded_file)
    
    # Display previously processed goals if available
    processed_goals = load_session_data('processed_goals')
    if processed_goals and not uploaded_file:
        st.markdown("---")
        st.markdown("### 📋 Previously Processed Goals")
        
        st.info("You have previously processed goals. You can review them below or upload a new document.")
        
        display_processed_goals(processed_goals)

def process_document(uploaded_file):
    """Process the uploaded document and extract goals"""
    
    st.markdown("---")
    st.markdown("### 🔍 Document Processing")
    
    # Update workflow status
    update_workflow_status('goals_upload', 'progress')
    
    with st.spinner("Processing document..."):
        # Initialize document processor
        doc_processor = DocumentProcessor()
        
        try:
            # Extract text from document
            extracted_text = doc_processor.extract_text(uploaded_file)
            
            if not extracted_text.strip():
                st.error("No text could be extracted from the document. Please check the file format.")
                return
            
            # Display extracted text preview
            st.markdown("#### 📄 Extracted Text Preview")
            with st.expander("View extracted text", expanded=False):
                st.text_area("Document content", extracted_text, height=200, disabled=True)
            
            # Process with AI agent
            process_with_ai_agent(extracted_text, uploaded_file.name)
            
        except Exception as e:
            st.error(f"Error processing document: {str(e)}")
            update_workflow_status('goals_upload', 'pending')

def process_with_ai_agent(text_content: str, filename: str):
    """Process the extracted text with Goal Validator Agent"""
    
    st.markdown("#### 🤖 AI Agent Analysis")
    
    # Initialize Goal Validator Agent
    goal_agent = GoalValidatorAgent()
    
    with st.spinner("AI agent is analyzing your goals..."):
        try:
            # Simulate agent processing time
            time.sleep(2)
            
            # Process goals with AI agent
            analysis_result = goal_agent.validate_goals(text_content)
            
            # Display results
            display_agent_analysis(analysis_result, filename)
            
        except Exception as e:
            st.error(f"Error during AI analysis: {str(e)}")
            update_workflow_status('goals_upload', 'pending')

def display_agent_analysis(analysis_result: Dict[str, Any], filename: str):
    """Display the results of AI agent analysis"""
    
    st.markdown("#### 📊 Analysis Results")
    
    # Overall assessment
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Goals Identified", analysis_result.get('goals_count', 0))
    
    with col2:
        smart_score = analysis_result.get('smart_score', 0)
        st.metric("SMART Score", f"{smart_score}%")
    
    with col3:
        quality_level = analysis_result.get('quality_level', 'Unknown')
        st.metric("Quality Level", quality_level)
    
    # Detailed goal analysis
    goals = analysis_result.get('goals', [])
    
    if goals:
        st.markdown("#### 🎯 Individual Goal Analysis")
        
        for i, goal in enumerate(goals, 1):
            with st.expander(f"Goal {i}: {goal.get('title', 'Untitled')}", expanded=True):
                
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    st.markdown("**Original Goal:**")
                    st.write(goal.get('original_text', ''))
                    
                    if goal.get('improved_version'):
                        st.markdown("**AI-Improved Version:**")
                        st.success(goal['improved_version'])
                
                with col2:
                    # SMART criteria assessment
                    st.markdown("**SMART Assessment:**")
                    smart_criteria = goal.get('smart_assessment', {})
                    
                    for criterion, status in smart_criteria.items():
                        if status:
                            st.success(f"✅ {criterion.title()}")
                        else:
                            st.error(f"❌ {criterion.title()}")
                    
                    # Issues and recommendations
                    if goal.get('issues'):
                        st.markdown("**Issues:**")
                        for issue in goal['issues']:
                            st.warning(f"⚠️ {issue}")
                    
                    if goal.get('recommendations'):
                        st.markdown("**Recommendations:**")
                        for rec in goal['recommendations']:
                            st.info(f"💡 {rec}")
        
        # Allow user to edit goals
        st.markdown("---")
        st.markdown("#### ✏️ Edit and Finalize Goals")
        
        edited_goals = edit_goals_interface(goals)
        
        # Save processed goals
        processed_data = {
            'filename': filename,
            'original_text': analysis_result.get('original_text', ''),
            'goals': edited_goals,
            'analysis_result': analysis_result,
            'processed_at': time.time()
        }
        
        save_session_data('processed_goals', processed_data)
        
        # Update session statistics
        if 'session_stats' not in st.session_state:
            st.session_state.session_stats = {
                'goals_processed': 0,
                'epics_generated': 0,
                'stories_analyzed': 0,
                'dependencies_found': 0
            }
        
        st.session_state.session_stats['goals_processed'] = len(edited_goals)
        
        # Mark step as complete
        update_workflow_status('goals_upload', 'complete')
        
        st.success("🎉 Goals processed and validated successfully!")
        
        # Next step guidance
        st.info("""
        **✅ Step 2 Complete!**
        
        Your PI goals have been validated and improved by AI agents.
        
        **Next Step:** Generate Epics and Features from your validated goals.
        """)
        
        if st.button("⚡ Continue to Generate Epics", use_container_width=True):
            st.switch_page("pages/3_⚡_Generate_Epics.py")
    
    else:
        st.warning("No clear goals were identified in the document. Please review the content and try again.")
        update_workflow_status('goals_upload', 'pending')

def edit_goals_interface(goals: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Provide interface for editing goals"""
    
    edited_goals = []
    
    st.info("Review and edit the AI-improved goals below. You can modify the text to better match your requirements.")
    
    for i, goal in enumerate(goals):
        st.markdown(f"**Goal {i + 1}:**")
        
        # Editable goal text
        improved_text = goal.get('improved_version', goal.get('original_text', ''))
        
        edited_text = st.text_area(
            f"Goal {i + 1} text:",
            value=improved_text,
            height=100,
            key=f"goal_edit_{i}"
        )
        
        # Additional metadata
        col1, col2 = st.columns(2)
        
        with col1:
            priority = st.selectbox(
                f"Priority for Goal {i + 1}:",
                options=['High', 'Medium', 'Low'],
                index=0,
                key=f"priority_{i}"
            )
        
        with col2:
            category = st.selectbox(
                f"Category for Goal {i + 1}:",
                options=['Business', 'Technical', 'User Experience', 'Performance', 'Security', 'Other'],
                index=0,
                key=f"category_{i}"
            )
        
        # Build edited goal
        edited_goal = {
            'title': goal.get('title', f'Goal {i + 1}'),
            'text': edited_text,
            'priority': priority,
            'category': category,
            'original_text': goal.get('original_text', ''),
            'smart_assessment': goal.get('smart_assessment', {}),
            'issues': goal.get('issues', []),
            'recommendations': goal.get('recommendations', [])
        }
        
        edited_goals.append(edited_goal)
        
        st.markdown("---")
    
    return edited_goals

def display_processed_goals(processed_data: Dict[str, Any]):
    """Display previously processed goals"""
    
    goals = processed_data.get('goals', [])
    filename = processed_data.get('filename', 'Unknown')
    
    st.info(f"**Source:** {filename}")
    
    # Summary metrics
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Goals", len(goals))
    
    with col2:
        high_priority = sum(1 for g in goals if g.get('priority') == 'High')
        st.metric("High Priority", high_priority)
    
    with col3:
        categories = set(g.get('category', 'Other') for g in goals)
        st.metric("Categories", len(categories))
    
    # Goal list
    for i, goal in enumerate(goals, 1):
        with st.expander(f"Goal {i}: {goal.get('title', 'Untitled')}", expanded=False):
            st.write(goal.get('text', ''))
            
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**Priority:** {goal.get('priority', 'Medium')}")
            with col2:
                st.write(f"**Category:** {goal.get('category', 'Other')}")
    
    # Action buttons
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("⚡ Continue to Generate Epics", use_container_width=True):
            st.switch_page("pages/3_⚡_Generate_Epics.py")
    
    with col2:
        if st.button("🔄 Upload New Document", use_container_width=True):
            # Clear processed goals
            if 'processed_goals' in st.session_state:
                del st.session_state['processed_goals']
            st.rerun()

def generate_demo_document(quality_type: str):
    """Generate a demo PI goals document using OpenAI"""
    
    config = load_config()
    openai_api_key = config.get('openai_api_key')
    
    if not openai_api_key:
        st.error("OpenAI API key not configured. Please set OPENAI_API_KEY in your environment.")
        return
    
    # Set up OpenAI client
    openai.api_key = openai_api_key
    
    st.markdown("---")
    st.markdown(f"### 🤖 Generating {quality_type.title()} Example Document")
    
    with st.spinner("AI is generating a sample PI goals document..."):
        try:
            if quality_type == "good":
                prompt = """
                Generate a comprehensive PI (Program Increment) goals document for a fictional e-commerce company's mobile app improvement initiative. 

                Create exactly 4 SMART goals using this format for each goal:

                GOAL 1: [Specific Action Verb] [What exactly will be accomplished]
                
                Success Criteria:
                - Achieve [specific number/percentage] improvement in [metric]
                - Complete implementation by [specific date within 12 weeks]
                - Deliver measurable business value of [specific amount/percentage]
                - Meet performance targets of [specific KPI numbers]
                
                Business Relevance: [How this directly impacts revenue/customers/efficiency]
                Timeline: Complete by end of PI (Week 12) with milestones at Week 4 and Week 8

                Example structure:
                GOAL 1: Implement mobile app checkout optimization system
                Success Criteria:
                - Reduce checkout abandonment rate by 25% (from current 40% to 30%)
                - Increase mobile conversion rate by 15% 
                - Complete development and testing by Week 10 of PI
                - Generate additional $500K monthly revenue through improved conversions
                Business Relevance: Directly improves customer experience and increases mobile revenue
                Timeline: Complete by Week 12 with beta testing in Week 8

                Make each goal follow this exact pattern with:
                - Specific action verbs (implement, develop, create, establish, build)
                - Exact percentages and numbers for all metrics
                - Clear deadlines within the 12-week PI timeframe
                - Quantified business impact
                - Realistic but ambitious targets

                Include 4 goals covering: performance optimization, user experience improvement, payment processing, and security enhancement.
                """
            else:  # poor quality
                prompt = """
                Generate a poorly written PI goals document for a fictional company that has common issues:
                
                Create 3-4 goals that have these problems:
                - Vague and non-specific language
                - No measurable metrics or unclear success criteria
                - Unrealistic timelines or no timelines at all
                - Missing business context
                - Unclear stakeholder responsibilities
                - No risk considerations
                - Ambiguous language like "improve", "enhance", "better"
                
                Make this document demonstrate typical problems that would fail SMART criteria validation.
                Include goals like "Make the app better", "Improve user experience", "Fix performance issues" without specifics.
                """
            
            # Call OpenAI API
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a business analyst creating PI planning documents. Generate realistic, detailed content."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            generated_content = response.choices[0].message.content
            
            # Display the generated document
            st.success(f"✅ Generated {quality_type} example document!")
            
            # Show document preview
            st.markdown("#### 📄 Generated Document Preview")
            with st.expander("View generated document", expanded=True):
                st.text_area("Generated content", generated_content, height=400, disabled=True)
            
            # Create Word document for download
            word_doc = create_word_document(generated_content, quality_type)
            
            # Download button
            st.markdown("#### 💾 Download Document")
            st.download_button(
                label=f"📥 Download {quality_type.title()} Example as Word Document",
                data=word_doc,
                file_name=f"PI_Goals_{quality_type.title()}_Example.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            st.info("""
            **💡 Next Steps:**
            1. Download the Word document above
            2. Upload it using the file uploader below to test the AI analysis
            3. See how the system validates and improves the goals
            """)
            
        except Exception as e:
            st.error(f"Error generating document: {str(e)}")
            st.info("Please check your OpenAI API key configuration.")

def create_word_document(content: str, quality_type: str) -> bytes:
    """Create a Word document from the generated content"""
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'PI Goals Document - {quality_type.title()} Example', 0)
    
    # Add subtitle
    subtitle = doc.add_heading('Generated by AI for Demonstration Purposes', level=2)
    
    # Add a line break
    doc.add_paragraph()
    
    # Split content into paragraphs and add to document
    paragraphs = content.split('\n\n')
    
    for paragraph in paragraphs:
        if paragraph.strip():
            # Check if it's a heading (starts with #, *, or is all caps)
            if (paragraph.strip().startswith('#') or 
                paragraph.strip().startswith('*') or 
                (paragraph.strip().isupper() and len(paragraph.strip()) < 100)):
                # Add as heading
                heading_text = paragraph.strip().lstrip('#*').strip()
                doc.add_heading(heading_text, level=1)
            else:
                # Add as regular paragraph
                doc.add_paragraph(paragraph.strip())
    
    # Add footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run("Generated by PI Planning Dashboard - ").italic = True
    footer_para.add_run(f"{quality_type.title()} Example Document").bold = True
    
    # Save to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer.getvalue()

def create_mock_file(content: str, filename: str):
    """Create a mock file object for processing generated content"""
    
    class MockFile:
        def __init__(self, content: str, name: str):
            self.content = content
            self.name = name
            self.size = len(content.encode('utf-8'))
        
        def read(self):
            return self.content.encode('utf-8')
        
        def getvalue(self):
            return self.content.encode('utf-8')
    
    return MockFile(content, filename)

if __name__ == "__main__":
    main()
